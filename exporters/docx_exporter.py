"""
docx_exporter - DOCX 导出器

直接从 TeacherRuntimePlan 生成 Word 文档。
不经过 Markdown 中转，不调用 LLM。

使用 python-docx 库。
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from models.runtime import TeacherRuntimePlan


def export_to_docx(plan: TeacherRuntimePlan, output_path: str) -> str:
    """
    将 TeacherRuntimePlan 直接导出为 DOCX 文件。

    Args:
        plan: 教师运行时教案
        output_path: 输出文件路径（如 outputs/lesson_plan.docx）

    Returns:
        实际写入的文件路径
    """
    doc = Document()

    # 设置默认字体
    _set_default_font(doc)

    # 1. 标题
    _add_title(doc, plan.topic)

    # 2. 基本信息表格
    _add_basic_info(doc, plan)

    # 3. 教学目标
    _add_objectives(doc, plan)

    # 4. 教学重点与难点
    _add_key_points(doc, plan)

    # 5. 教学过程（核心）
    _add_teaching_process(doc, plan)

    # 6. 课堂互动设计
    _add_interactions(doc, plan)

    # 7. 课堂练习
    _add_practice(doc, plan)

    # 8. 作业布置
    _add_homework(doc, plan)

    # 9. 板书设计
    _add_blackboard(doc, plan)

    # 10. 课堂小结
    _add_summary(doc, plan)

    doc.save(output_path)
    return output_path


# ============================================================
# 内部函数
# ============================================================

def _set_default_font(doc: Document):
    """设置文档默认字体"""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(11)
    # 设置中文字体
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _add_title(doc: Document, topic: str):
    """添加教案标题"""
    title = doc.add_heading(topic or "教学教案", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_basic_info(doc: Document, plan: TeacherRuntimePlan):
    """添加基本信息表格"""
    doc.add_heading("基本信息", level=1)

    table = doc.add_table(rows=0, cols=2, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    info_items = [
        ("课题", plan.topic),
        ("年级", plan.grade),
        ("课时", plan.duration),
    ]
    if plan.teaching_methods:
        info_items.append(("教学方法", "、".join(plan.teaching_methods)))

    for label, value in info_items:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
        # 加粗标签列
        for paragraph in row.cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True


def _add_objectives(doc: Document, plan: TeacherRuntimePlan):
    """添加教学目标"""
    if not plan.teaching_objectives:
        return

    doc.add_heading("教学目标", level=1)
    for i, obj in enumerate(plan.teaching_objectives, 1):
        doc.add_paragraph(f"{i}. {obj}")


def _add_key_points(doc: Document, plan: TeacherRuntimePlan):
    """添加教学重点与难点"""
    if not plan.key_points and not plan.difficult_points:
        return

    doc.add_heading("教学重点与难点", level=1)

    if plan.key_points:
        p = doc.add_paragraph()
        run = p.add_run("教学重点：")
        run.bold = True
        for point in plan.key_points:
            doc.add_paragraph(point, style="List Bullet")

    if plan.difficult_points:
        p = doc.add_paragraph()
        run = p.add_run("教学难点：")
        run.bold = True
        for point in plan.difficult_points:
            doc.add_paragraph(point, style="List Bullet")


def _add_teaching_process(doc: Document, plan: TeacherRuntimePlan):
    """添加教学过程（表格布局）"""
    if not plan.sections:
        return

    doc.add_heading("教学过程", level=1)

    for i, section in enumerate(plan.sections, 1):
        # 环节标题
        doc.add_heading(f"{i}. {section.title}", level=2)

        # 使用表格布局
        table = doc.add_table(rows=0, cols=2, style="Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        rows_data = [
            ("教师活动", section.teacher_activity),
            ("学生活动", section.student_activity),
        ]
        if section.interaction_method:
            rows_data.append(("互动方式", section.interaction_method))
        if section.duration_minutes:
            rows_data.append(("建议时长", f"{section.duration_minutes}分钟"))
        if section.teaching_intent:
            rows_data.append(("设计意图", section.teaching_intent))

        for label, value in rows_data:
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = value
            for paragraph in row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        doc.add_paragraph()  # 空行分隔


def _add_interactions(doc: Document, plan: TeacherRuntimePlan):
    """添加课堂互动设计"""
    if not plan.interactions:
        return

    doc.add_heading("课堂互动设计", level=1)

    for i, interaction in enumerate(plan.interactions, 1):
        doc.add_heading(f"互动{i}", level=2)

        table = doc.add_table(rows=0, cols=2, style="Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        rows_data = [
            ("触发时机", interaction.trigger),
            ("教师提问", interaction.teacher_question),
        ]

        if interaction.expected_responses:
            responses_text = "\n".join(f"- {r}" for r in interaction.expected_responses)
            rows_data.append(("预期学生回答", responses_text))

        if interaction.teacher_followup:
            rows_data.append(("教师追问", interaction.teacher_followup))

        for label, value in rows_data:
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = value
            for paragraph in row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        doc.add_paragraph()


def _add_practice(doc: Document, plan: TeacherRuntimePlan):
    """添加课堂练习"""
    if not plan.practice_questions:
        return

    doc.add_heading("课堂练习", level=1)

    table = doc.add_table(rows=1, cols=3, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    header_cells = table.rows[0].cells
    header_cells[0].text = "序号"
    header_cells[1].text = "题目"
    header_cells[2].text = "参考答案"
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # 题目
    for i, q in enumerate(plan.practice_questions, 1):
        row = table.add_row()
        row.cells[0].text = str(i)
        row.cells[1].text = f"[{q.difficulty}] {q.question}"
        row.cells[2].text = q.answer


def _add_homework(doc: Document, plan: TeacherRuntimePlan):
    """添加作业布置"""
    if not plan.homework:
        return

    doc.add_heading("作业布置", level=1)

    for hw in plan.homework:
        p = doc.add_paragraph()
        run = p.add_run(f"【{hw.type}】")
        run.bold = True
        p.add_run(f" {hw.content}")
        if hw.purpose:
            doc.add_paragraph(f"  目的：{hw.purpose}", style="List Bullet")


def _add_blackboard(doc: Document, plan: TeacherRuntimePlan):
    """添加板书设计"""
    if not plan.blackboard:
        return

    doc.add_heading("板书设计", level=1)

    bb = plan.blackboard

    if bb.layout:
        p = doc.add_paragraph()
        run = p.add_run("布局：")
        run.bold = True
        p.add_run(bb.layout)

    if bb.main_content:
        p = doc.add_paragraph()
        run = p.add_run("主板书：")
        run.bold = True
        for item in bb.main_content:
            doc.add_paragraph(item, style="List Bullet")

    if bb.key_formulas:
        p = doc.add_paragraph()
        run = p.add_run("核心公式/概念：")
        run.bold = True
        for item in bb.key_formulas:
            doc.add_paragraph(item, style="List Bullet")

    if bb.diagrams:
        p = doc.add_paragraph()
        run = p.add_run("图示：")
        run.bold = True
        for item in bb.diagrams:
            doc.add_paragraph(item, style="List Bullet")


def _add_summary(doc: Document, plan: TeacherRuntimePlan):
    """添加课堂小结"""
    if not plan.summary:
        return

    doc.add_heading("课堂小结", level=1)
    doc.add_paragraph(plan.summary)
